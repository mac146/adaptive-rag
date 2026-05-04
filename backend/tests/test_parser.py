from pathlib import Path

import docx

from app.ingestion.parser.parser import parse_docx, parse_markdown, parse_pdf


def test_pdf_extracts_text_from_sample():
    sample_path = Path(__file__).parent / "samples" / "test.pdf"

    result = parse_pdf(str(sample_path))

    assert result
    assert any(element["type"] in {"heading", "paragraph"} for element in result)
    assert any(element["text"].strip() for element in result)


def test_docx_extracts_headings_and_paragraphs(tmp_path):
    file_path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_heading("Project Overview", level=1)
    document.add_paragraph("This is a parser smoke test.")
    document.save(file_path)

    result = parse_docx(str(file_path))

    assert result == [
        {"type": "heading", "level": 1, "text": "Project Overview", "page": None},
        {"type": "paragraph", "level": None, "text": "This is a parser smoke test.", "page": None},
    ]


def test_docx_extracts_table_rows_with_section_metadata(tmp_path):
    file_path = tmp_path / "table.docx"
    document = docx.Document()
    document.add_heading("Timeline", level=1)
    table = document.add_table(rows=3, cols=3)

    headers = ["Stage", "Tool", "Date"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        run = cell.paragraphs[0].add_run(value)
        run.bold = True

    table.rows[1].cells[0].text = "Parsing"
    table.rows[1].cells[1].text = "python-docx"
    table.rows[1].cells[2].text = "2026-05-04"

    table.rows[2].cells[0].text = "Retrieval"
    table.rows[2].cells[1].text = "BM25"
    table.rows[2].cells[2].text = "2026-05-05"

    document.save(file_path)

    result = parse_docx(str(file_path))

    assert result == [
        {"type": "heading", "level": 1, "text": "Timeline", "page": None},
        {
            "type": "paragraph",
            "level": None,
            "text": "[Stage]: Parsing | [Tool]: python-docx | [Date]: 2026-05-04",
            "page": None,
            "section_title": "Timeline",
            "is_table_row": True,
            "table_index": 0,
            "row_index": 1,
        },
        {
            "type": "paragraph",
            "level": None,
            "text": "[Stage]: Retrieval | [Tool]: BM25 | [Date]: 2026-05-05",
            "page": None,
            "section_title": "Timeline",
            "is_table_row": True,
            "table_index": 0,
            "row_index": 2,
        },
    ]


def test_markdown_extracts_headings_and_paragraphs(tmp_path):
    file_path = tmp_path / "sample.md"
    file_path.write_text("# Title\n\nIntro paragraph.\n\n## Details\n\nMore context.\n", encoding="utf-8")

    result = parse_markdown(str(file_path))

    assert result == [
        {"type": "heading", "level": 1, "text": "Title", "page": None},
        {"type": "paragraph", "level": None, "text": "Intro paragraph.", "page": None},
        {"type": "heading", "level": 2, "text": "Details", "page": None},
        {"type": "paragraph", "level": None, "text": "More context.", "page": None},
    ]
